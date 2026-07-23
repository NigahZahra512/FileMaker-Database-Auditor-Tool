"""
docx_report.py

GROUP A FEATURE: DOCX Export
--------------------------------
Inspired by the "Download DOCX" button on the FM Changelog reference
tool's Deep Audit screen. Takes the same report dict every endpoint in
main.py already returns ({"summary": {...}, "findings": [...]}) and
builds a Word document out of it, so a non-technical client can open
the audit results in Word instead of a browser.

Mirrors the frontend's own grouping logic (frontend/index.html's
categorySectionsHtml): if a finding has a "category", group by that;
DDR results fall under Fields / Scripts / Relationships / Layouts /
Unused Fields / Unused Scripts / Call Chain, in that order. Findings
from the Script Review or SQL Review tabs don't carry a "category" --
those fall back to their "module" (e.g. "script", "sql") as a single
group, matching how renderReport() shows them on the non-DDR tabs.

Nothing here talks to the AI providers or touches the filesystem on
disk -- build_docx_report() returns an in-memory BytesIO buffer, and
main.py streams that straight back as the HTTP response body. Keeps
the same "never write anything the user didn't ask to download"
pattern as the rest of the backend.
"""

import io
from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

SEVERITY_ORDER = {"Critical": 0, "Warning": 1, "Info": 2}
CATEGORY_ORDER = [
    "Fields", "Scripts", "Relationships", "Layouts",
    "Unused Fields", "Unused Scripts", "Call Chain",
]

SEVERITY_COLORS = {
    "Critical": RGBColor(0xB0, 0x20, 0x20),
    "Warning": RGBColor(0xB0, 0x7A, 0x10),
    "Info": RGBColor(0x30, 0x60, 0xA0),
}


def _group_findings(findings):
    groups = {}
    for f in findings:
        cat = f.get("category") or f.get("module") or "Findings"
        groups.setdefault(cat, []).append(f)
    ordered_keys = (
        [c for c in CATEGORY_ORDER if c in groups]
        + [c for c in groups if c not in CATEGORY_ORDER]
    )
    return [(cat, groups[cat]) for cat in ordered_keys]


def _add_summary_table(doc, summary):
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"
    headers = ["Critical", "Warning", "Info"]
    values = [summary.get("critical", 0), summary.get("warning", 0), summary.get("info", 0)]
    for col, (label, value) in enumerate(zip(headers, values)):
        head_cell = table.cell(0, col)
        head_cell.text = label
        head_cell.paragraphs[0].runs[0].bold = True
        table.cell(1, col).text = str(value)
        table.cell(1, col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_findings_table(doc, items):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for cell, label in zip(hdr, ["Severity", "Location", "Description", "Suggestion"]):
        cell.text = label
        cell.paragraphs[0].runs[0].bold = True

    for f in items:
        row = table.add_row().cells
        sev = f.get("severity", "Info")
        sev_run = row[0].paragraphs[0].add_run(sev)
        sev_run.bold = True
        sev_run.font.color.rgb = SEVERITY_COLORS.get(sev, RGBColor(0, 0, 0))
        row[1].text = f.get("location", "") or ""
        row[2].text = f.get("description", "") or ""
        row[3].text = f.get("suggestion", "") or ""


def _add_stats_table(doc, pairs):
    """pairs: list of (label, value) -- rendered as a compact grid,
    same idea as _add_summary_table but for an arbitrary number of
    stat cells instead of always exactly three."""
    cols = 4
    rows_needed = -(-len(pairs) // cols)  # ceil
    table = doc.add_table(rows=rows_needed * 2, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(pairs):
        r, c = divmod(i, cols)
        head_cell = table.cell(r * 2, c)
        head_cell.text = label
        head_cell.paragraphs[0].runs[0].bold = True
        val_cell = table.cell(r * 2 + 1, c)
        val_cell.text = str(value)
        val_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_table_audit_docx(detail: dict, snapshot_label: str) -> io.BytesIO:
    """Deep Audit export for a single table -- the DOCX counterpart of
    the Explore tab's "deep audit" link. Takes the same detail dict
    table_audit.build_table_detail() returns (fields, relationships,
    layouts, scripts, plus the auto-enter / dynamic-access / typo
    detections) and lays it out as a standalone Word document,
    mirroring the reference tool's per-table Deep Audit screen."""
    fields = detail.get("fields", []) or []
    relationships = detail.get("relationships", []) or []
    layouts = detail.get("layouts", []) or []
    scripts = detail.get("scripts", []) or []
    other_auto_enter_calcs = detail.get("other_auto_enter_calcs", []) or []
    dynamic_access_fields = detail.get("dynamic_access_fields", []) or []
    likely_typos = detail.get("likely_typos", []) or []

    unstored_count = sum(1 for f in fields if f.get("is_calculation") and f.get("storage_index") == "None")
    always_eval_count = sum(1 for f in fields if f.get("always_evaluate"))
    unused_count = sum(1 for f in fields if f.get("is_unused"))
    validated_count = sum(1 for f in fields if f.get("has_validation"))
    global_count = sum(1 for f in fields if f.get("is_global"))

    doc = Document()

    title = doc.add_heading(f'{detail.get("table_name", "")} \u2014 Deep Audit', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("Table deep audit \u00b7 DDR snapshot").italic = True

    meta = doc.add_paragraph()
    meta.add_run(f"Source: {snapshot_label}\n").italic = True
    meta.add_run(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    ).italic = True

    if always_eval_count == 0 and not dynamic_access_fields:
        doc.add_paragraph("No alwaysEvaluate or dynamic access issues \u2014 see table summary below.")
    else:
        doc.add_paragraph(
            f"{always_eval_count} always-evaluate field(s) and {len(dynamic_access_fields)} "
            "dynamic-access field(s) found \u2014 review before making changes."
        )

    doc.add_heading("Summary", level=1)
    _add_stats_table(doc, [
        ("Total Fields", len(fields)),
        ("Records", detail.get("record_count", 0)),
        ("AlwaysEval AE", always_eval_count),
        ("Auto-Enter Total", detail.get("auto_enter_total", 0)),
        ("Calc Fields", detail.get("calc_field_count", 0)),
        ("Unstored Calcs", unstored_count),
        ("Dynamic Access", len(dynamic_access_fields)),
        ("Globals", global_count),
        ("Likely Typos", len(likely_typos)),
        ("Unused Fields", unused_count),
        ("Validated Fields", validated_count),
    ])

    doc.add_heading(f"Fields ({len(fields)})", level=1)
    if fields:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for cell, label in zip(table.rows[0].cells, ["Name", "Type", "Kind", "Flags"]):
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True
        for f in fields:
            flags = []
            if f.get("is_calculation") and f.get("storage_index") == "None":
                flags.append("Unstored")
            if f.get("always_evaluate"):
                flags.append("Always-eval")
            if f.get("is_unused"):
                flags.append("Unused")
            if f.get("has_validation"):
                flags.append("Validated")
            if f.get("is_global"):
                flags.append("Global")
            row = table.add_row().cells
            row[0].text = f.get("name", "") or ""
            row[1].text = f.get("data_type", "") or ""
            row[2].text = "Calculation" if f.get("is_calculation") else (f.get("field_type") or "Normal")
            row[3].text = ", ".join(flags) or "\u2014"
    else:
        doc.add_paragraph("No fields on this table.")

    doc.add_heading(f"Other Auto-Enter Calcs ({len(other_auto_enter_calcs)})", level=1)
    doc.add_paragraph(
        "Fields with auto-enter calculations that fire on record creation or when the "
        "field is empty, but do not re-evaluate on every commit."
    ).italic = True
    if other_auto_enter_calcs:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        for cell, label in zip(table.rows[0].cells, ["Field", "Type", "Calc"]):
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True
        for f in other_auto_enter_calcs:
            row = table.add_row().cells
            row[0].text = f.get("name", "") or ""
            row[1].text = f.get("data_type", "") or ""
            row[2].text = f.get("calculation_text", "") or ""
    else:
        doc.add_paragraph("No non-always-evaluate auto-enter calcs on this table.")

    doc.add_heading(f"Dynamic Access ({len(dynamic_access_fields)})", level=1)
    if dynamic_access_fields:
        for d in dynamic_access_fields:
            doc.add_paragraph(f'{d.get("name", "")} \u2014 uses {d.get("function", "")}()')
    else:
        doc.add_paragraph("No fields reach other data through ExecuteSQL, GetField, or Evaluate.")

    doc.add_heading(f"Likely Typos ({len(likely_typos)})", level=1)
    if likely_typos:
        for t in likely_typos:
            doc.add_paragraph(f'{t.get("field_a", "")} \u2194 {t.get("field_b", "")} (edit distance {t.get("distance", "")})')
    else:
        doc.add_paragraph("No suspiciously similar field name pairs found.")

    doc.add_heading(f"Related Tables ({len(relationships)})", level=1)
    if relationships:
        for r in relationships:
            preds = ", ".join(
                f'{(p.get("left_field") or {}).get("name", "?")} {p.get("type", "=")} '
                f'{(p.get("right_field") or {}).get("name", "?")}'
                for p in r.get("predicates", [])
            )
            doc.add_paragraph(f'{r.get("other_table", "")} \u2014 {preds}')
    else:
        doc.add_paragraph("No relationships involve this table.")

    doc.add_heading(f"Layouts Using This Table ({len(layouts)})", level=1)
    if layouts:
        for l in layouts:
            fields_used = l.get("fields_used", [])
            doc.add_paragraph(f'{l.get("layout_name", "")} \u2014 {len(fields_used)} field(s) used')
    else:
        doc.add_paragraph("No layouts place a field from this table.")

    doc.add_heading(f"Scripts Touching This Table ({len(scripts)})", level=1)
    if scripts:
        for s in scripts:
            touched = s.get("fields_touched", [])
            doc.add_paragraph(f'{s.get("script_name", "")} \u2014 touches {len(touched)} field(s)')
    else:
        doc.add_paragraph("No script step directly references a field from this table.")

    doc.add_heading("Verification Before Deploy", level=1)
    for note in (
        "Before deleting any field, verify it is not referenced by external systems "
        "(API clients, data exports, webhook payloads, sync jobs). Static analysis "
        "cannot see consumers that live outside the DDR.",
        "For any field with dynamic access (ExecuteSQL, GetField, Evaluate), the list "
        "above is a best-effort extraction of the compile-time literal cases. Walk each "
        "possible runtime argument and confirm the full target set before changing the field.",
        "Always verify output against a representative sample of historical records and "
        "re-run the downstream exports and reports that read these fields \u2014 the "
        "before/after diff must be zero before deploying the change.",
    ):
        doc.add_paragraph(note, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_docx_report(report: dict, source_label: str) -> io.BytesIO:
    """Builds a Word document from a report dict and returns it as an
    in-memory BytesIO buffer, positioned at the start and ready to be
    streamed back as an HTTP response."""
    findings = report.get("findings", []) or []
    summary = report.get("summary", {}) or {}
    findings = sorted(findings, key=lambda f: SEVERITY_ORDER.get(f.get("severity"), 99))

    doc = Document()

    title = doc.add_heading("FileMaker Database Audit Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run(f"Source: {source_label}\n").italic = True
    meta.add_run(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    ).italic = True

    doc.add_heading("Summary", level=1)
    if findings:
        _add_summary_table(doc, summary)
    else:
        doc.add_paragraph("No findings -- looks clean, or nothing has been analysed yet.")

    for category, items in _group_findings(findings):
        doc.add_heading(category, level=1)
        doc.add_paragraph(f"{len(items)} finding{'s' if len(items) != 1 else ''}")
        _add_findings_table(doc, items)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
